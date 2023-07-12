from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import Paragraph
import docx


def dummy_pdf(text: str, temp_file_path="/tmp/temp_file"):
    c = canvas.Canvas(temp_file_path, pagesize=letter)

    style = ParagraphStyle(
        name="Normal", alignment=TA_CENTER, fontName="Helvetica", fontSize=12
    )

    p = Paragraph(text, style=style)
    p.wrapOn(c, inch * 6, inch * 4)
    p.drawOn(c, inch * 2, inch * 5)

    c.save()

    return temp_file_path


def dummy_txt(text: str):
    temp_file_path = "/tmp/temp_file"

    with open(temp_file_path, "w") as file:
        file.write(text)

    return temp_file_path


def dummy_docx(text: str):
    temp_file_path = "/tmp/temp_file"

    document = docx.Document()

    document.add_paragraph(text)

    # add a table to the document
    table = document.add_table(rows=3, cols=3)
    for i in range(3):
        row_cells = table.rows[i].cells
        for j in range(3):
            row_cells[j].text = f"Row {i+1}, Column {j+1}"

    document.save(temp_file_path)

    return temp_file_path
